import streamlit as st
from langchain_openai import ChatOpenAI
#from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain, SimpleSequentialChain
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import os
import pdfplumber
import pandas as pd
import extract_msg
import re
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from langchain_openai import AzureChatOpenAI
from openpyxl import load_workbook
import json

def expand_product_categories(impacted_products_text, product_alignment):
    if not product_alignment or not impacted_products_text:
        return impacted_products_text
    
    def extract_impact_status_from_table(text):
        impact_status = {}
        lines = text.split('\n')
        
        # More specific indicators
        positive_indicators = ['yes', 'y', 'true', '1', 'impacted', 'affected']
        negative_indicators = ['no', 'n', 'false', '0', 'not impacted', 'not affected', 'na', 'n/a']
        
        for line in lines:
            if '---' in line or '===' in line:
                continue
                
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')]
                cells = [cell for cell in cells if cell]
                
                if len(cells) >= 2:
                    category_cell = cells[0].lower().strip()
                    
                    # CRITICAL: Skip "All" or "ALL" completely
                    if category_cell in ['all', 'all products', 'all categories']:
                        continue
                    
                    # Check for exact matches with JSON keys
                    matched_category = None
                    for json_key in product_alignment.keys():
                        if json_key.lower() == category_cell:
                            matched_category = json_key
                            break
                        elif json_key.lower() in category_cell:
                            matched_category = json_key
                            break
                        elif category_cell == 'endowment' and json_key == 'endowment_plans':
                            matched_category = json_key
                            break
                    
                    if matched_category:
                        # Check status in subsequent cells
                        category_status = False
                        for status_cell in cells[1:]:
                            status_lower = status_cell.lower().strip()
                            if any(indicator in status_lower for indicator in positive_indicators):
                                category_status = True
                                break
                            elif any(indicator in status_lower for indicator in negative_indicators):
                                category_status = False
                                break
                        
                        impact_status[matched_category] = category_status
        
        return impact_status
    
    # Extract impact status
    impact_status = extract_impact_status_from_table(impacted_products_text)
    
    # Sanitize the "### 2.1 Impacted Products" section: keep only the first markdown table, drop any lists/headings that LLM may have added
    try:
        lower_text = impacted_products_text.lower()
        start_tokens = ["### 2.1 impacted products", "## 2.1 impacted products"]
        end_tokens = ["### 2.2", "## 2.2", "### 2.2 applications impacted", "## 2.2 applications impacted"]
        start_idx = -1
        for t in start_tokens:
            si = lower_text.find(t)
            if si != -1:
                start_idx = si
                break
        if start_idx != -1:
            end_idx = len(impacted_products_text)
            for t in end_tokens:
                ei = lower_text.find(t, start_idx + 1)
                if ei != -1:
                    end_idx = min(end_idx, ei)
            section = impacted_products_text[start_idx:end_idx]
            section_lines = section.split('\n')
            kept = []
            table_started = False
            table_ended = False
            for ln in section_lines:
                if not table_started:
                    kept.append(ln)
                    if '|' in ln:
                        table_started = True
                else:
                    if ('|' in ln) and not table_ended:
                        kept.append(ln)
                    else:
                        # once a non-table line appears after table has started, stop keeping further lines
                        table_ended = True
                
            sanitized_section = '\n'.join(kept)
            impacted_products_text = impacted_products_text[:start_idx] + sanitized_section + impacted_products_text[end_idx:]
    except Exception:
        pass
    
    # Only expand categories with explicit "Yes" status
    expanded_sections = []
    for category, products in product_alignment.items():
        if products and impact_status.get(category, False):  # Only if explicitly True
            if impact_status[category] == "Yes" or impact_status[category] == "All":
                product_list = '\n'.join([f"  - {product}" for product in products])
                category_display = category.upper().replace('_', ' ')
                category_section = f"\n\n**{category_display} Products (Impacted - Yes):**\n{product_list}"
                expanded_sections.append(category_section)
    
    # Append expansions to sanitized text
    if expanded_sections:
        return impacted_products_text + ''.join(expanded_sections)
    else:
        return impacted_products_text

def load_product_alignment():

    try:
        product_alignment = {
            "annuity": [
                "Bajaj Allianz Life Guaranteed Pension Goal II",
                "Bajaj Allianz Life Saral Pension"
            ],
            "combi": [
                "Bajaj Allianz Life Capital Goal Suraksha"
            ],
            "group": [
                "Bajaj Allianz Life Group Term Life",
                "Bajaj Allianz Life Group Credit Protection Plus",
                "Bajaj Allianz Life Group Sampoorna Jeevan Suraksha",
                "Bajaj Allianz Life Group Employee Benefit",
                "Bajaj Allianz Life Group Superannuation Secure Plus",
                "Bajaj Allianz Life Group Superannuation Secure",
                "Bajaj Allianz Life Group Employee Care",
                "Bajaj Allianz Life Group Secure Return",
                "Bajaj Allianz Life Group Sampoorna Suraksha Kavach",
                "Bajaj Allianz Life Pradhan Mantri Jeevan Jyoti Bima Yojana",
                "Bajaj Allianz Life Group Secure Shield",
                "Bajaj Allianz Life Group Investment Plan"
            ],
            "non_par": [
                "Bajaj Allianz Life Goal Suraksha",
                "Bajaj Allianz Life Assured Wealth Goal Platinum",
                "Bajaj Allianz Life Guaranteed Wealth Goal",
                "Bajaj Allianz Life Guaranteed Saving Goal",
                "Bajaj Allianz Life Assured Wealth Goal"
            ],
            "par": [
                "Bajaj Allianz Life ACE",
                "Bajaj Allianz Life ACE Advantage"
            ],
            "rider": [
                "Bajaj Allianz Accidental Death Benefit Rider",
                "Bajaj Allianz Accidental Permanent Total/Partial Disability Benefit Rider",
                "Bajaj Allianz Life Linked Accident Protection Rider II",
                "Bajaj Allianz Life Family Protect Rider",
                "Bajaj Allianz Life Group New Terminal Illness Rider",
                "Bajaj Allianz Life Group Accelerated Critical Illness Rider",
                "Bajaj Allianz Life Group Accidental Permanent Total/Partial Disability Benefit Rider",
                "Bajaj Allianz Life Group Critical Illness Rider",
                "Bajaj Allianz Life Group Accidental Death Benefit",
                "Bajaj Allianz Life New Critical Illness Benefit Rider",
                "Bajaj Allianz Life Care Plus Rider",
                "Bajaj Allianz Life Linked Critical Illness Benefit Rider"
            ],
            "term": [
                "Bajaj Allianz Life iSecure II",
                "Bajaj Allianz Life eTouch II",
                "Bajaj Allianz Life Saral Jeevan Bima",
                "Bajaj Allianz Life Diabetic Term Plan II Sub 8 HbA1c",
                "Bajaj Allianz Life Smart Protection Goal"
            ],
            "ulip": [
                "Bajaj Allianz Life Goal Assure IV",
                "Bajaj Allianz Life Magnum Fortune Plus III",
                "Bajaj Allianz Life Invest Protect Goal III",
                "Bajaj Allianz Life Fortune Gain II",
                "Bajaj Allianz Life Future Wealth Gain IV",
                "Bajaj Allianz Life LongLife Goal III",
                "Bajaj Allianz Life Smart Wealth Goal V",
                "Bajaj Allianz Life Goal Based Saving III",
                "Bajaj Allianz Life Elite Assure"
            ],
            "ulip_pension": [
                "Bajaj Allianz Life Smart Pension"
            ],
            "endowment_plans": [
                "Bajaj Allianz Life Assured Wealth Goal Platinum",
                "Bajaj Allianz Life ACE",
                "Bajaj Allianz Life Goal Suraksha"
            ]
        }
        
        return product_alignment
        
    except Exception as e:
        print(f"Error loading product alignment: {str(e)}")
        return {}

BRD_FORMAT = """
## 1.0 Introduction
    ## 1.1 Purpose
    ## 1.2 As-is process
    ## 1.3 To be process / High level solution
## 2.0 Impact Analysis
    ## 2.1 Impacted Products
    ## 2.2 Applications Impacted
    ## 2.3 List of APIs required
## 3.0 Process / Data Flow diagram / Figma
## 4.0 Business / System Requirement
## 5.0 MIS / DATA Requirement
## 6.0 Communication Requirement
## 7.0 Test Scenarios
## 8.0 Questions / Suggestions
## 9.0 Reference Document
## 10.0 Appendix
## 11.0 Risk Evaluation
"""

SECTION_TEMPLATES = {
 
    "intro_impact": """
 
You are a Business Analyst expert creating sections 1.0–2.0 of a comprehensive Business Requirements Document (BRD).
 
IMPORTANT: Do not output any ``` code fences or Mermaid syntax.
All text should be plain markdown (headings, lists, tables) only - no code blocks or fenced content.
Never expose the processing steps and instructions to the user while creating the BRD.
- Do not include "\n", "\\n", "/n", "<br>", "<br/>", or any other escape/HTML line break.
- For new line, just insert an actual line break (press Enter).
- For paragraph break, insert one blank line (double Enter).
 
SOURCE REQUIREMENTS:
 
{requirements}
 
EXCEL FILE PROCESSING INSTRUCTIONS:
 
**FOR EXCEL FILES (.xlsx/.xls):**
- Process ALL sheets EXCEPT "Test Scenarios" sheet
- **PRIORITY FOCUS**: Look specifically for "PART B : (Mandatory) Detailed Requirement" section in any sheet
- Include data from sheets: "Requirement", "Ops Risk Assessment", and any other available sheets
- Extract content from ALL relevant columns and rows in each sheet
- Look for business requirements, processes, impacts, and technical specifications across all sheets
- If sheet names are different from expected, process all sheets except those explicitly containing test scenarios
 
**SPECIAL INSTRUCTION FOR PURPOSE AND TO-BE PROCESS:**
For sections 1.1 Purpose and 1.3 To be process / High level solution:
- **PRIMARY PRIORITY**: Search for and extract information from "PART B : (Mandatory) Detailed Requirement" section
- Look for this exact text or similar variations like:
  - "PART B"
  - "Mandatory Detailed Requirement"
  - "Detailed Requirement"
  - "Part B - Detailed Requirement"
  - "PART B : Detailed Requirement"
- If found, prioritize this section's content for Purpose and To-be process extraction
- If not found, then search across ALL other processed sheets for relevant content
 
CRITICAL INSTRUCTIONS:
 
- Extract information from ALL available sheets (except Test Scenarios sheet)
- **For Purpose and To-be process: PRIORITIZE "PART B"  Detailed Requirement" content**
- Identify the ACTUAL business problem being solved from any relevant sheet
- Focus on what is explicitly mentioned across all processed sheets
- Do NOT create, assume, or fabricate any content not present in the source
- If a section has no relevant information across ALL processed sheets, leave it BLANK
- Adapt to any domain (training, payments, integration, access control, etc.)
 
Create ONLY the following sections with detailed content in markdown:
 
## 1.0 Introduction
 
### 1.1 Purpose
 
**SEARCH STRATEGY FOR PURPOSE:**
1. **FIRST PRIORITY**: Look specifically for "PART B : (Mandatory) Detailed Requirement" section
2. **SECOND PRIORITY**: Search other sections in "Requirement" sheet
 
Extract the EXACT business purpose, focusing on:
- Capture from "PART B : (Mandatory) Detailed Requirement" if available
- If PART B not found, capture from "Detailed Requirement" sections in any sheet
- What is the main business objective or problem being addressed?
- What specific functionality or capability is being implemented?
- What restrictions, validations, or controls are being introduced?
- What business processes are being improved or changed?
- What compliance, security, or operational requirements are being met?
 
Search across ALL processed sheets for key phrases: "purpose", "objective", "requirement", "need", "problem", "solution", "implement", "restrict", "validate", "improve", "ensure"
 
**EXTRACTION PRIORITY ORDER:**
1. Content from "PART B : (Mandatory) Detailed Requirement"
2. Content from other "Detailed Requirement" sections
3. Content from other relevant sections across all sheets
 
**CRITICAL**:
1. Do not use bullet points, numbering, or line breaks inside Purpose for output
2. If multiple lines exist, merge them into one cohesive paragraph with proper sentence flow.
 
### 1.2 As-is process
 
**FORMAT: Present content as BULLET POINTS using markdown bullet format (- or *)**
 
Extract the CURRENT state/process from ANY relevant sheet:
- How does the current system/process work?
- What are the existing workflows or user journeys?
- What problems or limitations exist in the current approach?
- What manual processes or workarounds are currently used?
- What system behaviors need to be changed?
- Any screenshots, process flows, or current state descriptions
 
Look for indicators across ALL sheets: "currently", "as-is", "existing", "present", "manual", "workaround", "problem with current", "limitations"
 
**CRITICAL: Format ALL extracted content as bullet points (- or *) - DO NOT use paragraphs or numbered lists**
 
### 1.3 To be process / High level solution
 
**FORMAT: Present content as BULLET POINTS using markdown bullet format (- or *)**
 
**SEARCH STRATEGY FOR TO-BE PROCESS:**
1. **FIRST PRIORITY**: Look specifically for "PART B : (Mandatory) Detailed Requirement" section
2. **SECOND PRIORITY**: Search other sections in "Requirement" sheet
3. **THIRD PRIORITY**: Search "Ops Risk Assessment" and other sheets
 
Extract the PROPOSED solution, prioritizing "PART B : (Mandatory) Detailed Requirement" content:
- Content from "PART B : (Mandatory) Detailed Requirement" if available
- What is the new process or system behavior?
- What workflow steps or validation logic will be implemented?
- How will the new solution address current problems?
- What automated processes will replace manual ones?
- What new capabilities or features will be added?
- Any conditional logic, decision trees, or multi-step processes
 
Look for indicators across ALL sheets: "to-be", "proposed", "solution", "new process", "will be", "should be", "automated", "enhanced", "improved", "step-by-step", "workflow", "condition", "if-then"
 
**CRITICAL: Format ALL extracted content as bullet points (- or *) - DO NOT use paragraphs or numbered lists**
 
**EXTRACTION PRIORITY ORDER:**
1. Content from "PART B : (Mandatory) Detailed Requirement"
2. Content from other "Detailed Requirement" sections
3. Content from other relevant sections across all sheets
 
 
## 2.0 Impact Analysis
 
### 2.1 Impacted Products
 
STEP BY STEP Process:
1. From part_c, extract the list `data_rows` containing "Type of Product" and their "List of products in which the change has to be done", For Example.,  
   ```json
        "data_rows": [
           curly bracket
                "row_description": "List of products in which the change has to be done",
                "values": curly bracket
                    "Type of Product": "List of products in which the change has to be done",
                    "ULIP": "-",
                    "TERM": "-",
                    "All": "Yes"
                curly bracket
            curly bracket
        ]
2. Create a list of all product names with their impact status.
 
3. From the section '=== PRODUCT ALIGNMENT DATA ===' in the source requirements,
   identify all the product names from Step 2 that EXACTLY match the JSON keys in PRODUCT ALIGNMENT DATA.
 
   - If no any product names match, STOP here and output:
     No impacted products found.
 
   - Do NOT proceed to impact status check unless a match with PRODUCT ALIGNMENT DATA keys is found.and Donot stop only on first match ,there can be multiple product which can match with PRODUCT ALIGNMENT DATA keys.
 
4. For every matched product name, apply IMPACT STATUS check as a mandatory second filter:
   - A product is eligible for expansion ONLY IF:
        (product_name is present in PRODUCT ALIGNMENT DATA keys) AND (impact_status is "Yes" OR "All")
   - If impact_status is "-" or "No" or "NA" or blank → DO NOT EXPAND, even if the product_name matches.
   - This rule is absolute. Example: If Prduct exists in PRODUCT ALIGNMENT DATA but its status is "-",
     then Product must NOT be expanded and should be excluded from the final output.
5. Expansion Rule:
   - For every product name that passes Step 4 (status is "Yes" or "All"), expand its entire mapped product list from PRODUCT ALIGNMENT DATA.
   -**VERY CRITICAL** If multiple product names qualify, expand all of them, not just the first.
   - The final output table must include every qualifying product category and each of its mapped values as separate rows.
   - Do not skip or collapse duplicates. Every eligible mapping must appear in the output table.
6. Important :Format the output as a markdown table with the following structure in product category all product names which get qualified and it a product should be repeated util it shows all its mapped product values:
 
| Product Category | Individual Products Name |
|------------------|---------------------------|
| [PRODUCT_NAME1]   | [Product 1]              |
| [PRODUCT_NAME1]   | [Product 2]              |
| [PRODUCT_NAME2]   | [Product 3]              |
| [PRODUCT_NAME2]   | [Product 4]              |
| [PRODUCT_NAME3]   | [Product 5]              |
| [PRODUCT_NAME3]   | [Product 6]              |
 
 
7. If the final filtered list is empty (i.e., no products with status "Yes" or "All"), output exactly:
    No impacted products found.
 
---
 
### VERY IMPORTANT NOTES:
- Do NOT expand or include any product whose status is not explicitly "Yes" or "All".
- Do NOT assume impact; strictly follow the source status.
- Maintain the exact markdown format above.
 
---
 
 
 
VERY VERY CRITICAL  VALIDATION RULES:
 
1. Matching with PRODUCT ALIGNMENT DATA keys:
   - Perform **case-insensitive exact key match**.
   - Example: "term" in part_c matches "TERM" in PRODUCT ALIGNMENT DATA.
 
2. Impact status normalization:
   - Convert all statuses to lowercase before comparison.
   - Treat "yes", "all" (in any case: Yes/YES/All/ALL) as positive.
   - Treat "-", "no", "na", "" (blank) as negative.
 
3. Multi-product expansion:
   - If multiple products pass the filter, expand **ALL of them**.
   - Never stop after the first match. Iterate through all qualifying products.
 
4. No fallbacks:
   - Do not assume impact if not matched.
   - Do not collapse duplicates.
 
### 2.2 Applications Impacted
 
STEP BY STEP Process:
1. From part_c, extract the list `data_rows` containing "Application Name" and their "Pls select correct response", e.g.,  
   ```json
        "data_rows": [
           curly bracket
                "row_description": "Pls select correct response",
                "values": curly bracket
                    "Type of Product": "Pls select correct response",
                    "OPUS": "-",
                    "INSTAB": "-",
                    "Other": "DigiAgency"
                curly bracket
            curly bracket
        ]
 
2. Extract applications list from part_c (Application Name : Pls select correct response).
 
3.Filter logic:
 
    If value = "-", "No", "NA", or "" (blank) → exclude from output completely.
 
    Otherwise (any other value) → include.
 
4.Special case for "Other":
 
    If "Other" has a valid value (after filter), replace "Other" with that value as the Application Name.
 
5. For every application that passed the filter, output in markdown table containing two columns:
   - **Application Name**
   - **High level Description**: a short 1–2 line explanation of how this application is impacted by the change.
     IMPORTANT: Do not copy the placeholder text "High level descriptions of Applications basically the overview how it is impacted".
     Instead, generate a meaningful description based on the application name and context.
| Application Name | High level Description |
| DigiAgency | Impact description of how App is impacted |
 
**VALIDATION RULE:**
- List ONLY the applications explicitly with an impact status of those application whose vakue of Pls select correct response is anything except for "-", "" and "No","NA" and "".
 
### 2.3 List of APIs required
 
Extract SPECIFIC technical requirements from ALL processed sheets:
- New APIs or services that need to be created
- Existing APIs that need modification
- Third-party integrations or external system connections
- Database access or query requirements
- Authentication, authorization, or security services
- Any technical specifications or interface requirements
 
CATALOG MATCHING (use appended block titled "=== KNOWN API CATALOG (READ-ONLY REFERENCE) ==="):
- Parse the JSON catalog provided in the appended block
- For each requirement, first attempt to MATCH the requirement description/intent with catalog descriptions
- If a match is found, output the EXACT method and endpoint from the catalog (do not modify), and use the catalog description
- If a requirement has no catalog match, add a "Custom API – [METHOD] [endpoint]" row with a concise description from the source requirement
 
OUTPUT FORMAT (MANDATORY TABLE ONLY):
| S. No | API Name | API Description |
|-------|----------|-----------------|
| 1 | GET /AgentDetails | Retrieve agent details including training completion status. |
| 2 | POST /TrainingStatusValidation | Validate training completion flag. |
| 3 | Custom API – POST /DisplayRestrictionMessage | Display restriction message if training is incomplete. |
 
HARD RULES:
- OUTPUT ONLY the table above (no paragraphs, bullet lists, or extra text before/after)
- Table must have EXACTLY 3 columns: S. No, API Name, API Description.
- API Name must be exactly "[METHOD] [endpoint]" for catalog matches
- Do NOT invent/alter catalog endpoints or methods.
- Do NOT add extra columns or split descriptions across multiple columns
- If nothing is identified, output a single-row table stating "No APIs identified from source"
 
IMPORTANT:
 
- Use markdown headings (##, ###)
- **CRITICAL**: For sections 2.1 and 2.2, if structured tables exist in source, reproduce them as markdown tables
- Sections 1.2 and 1.3 to be in form of bullet pointers
- **For Purpose and To-be process: PRIORITIZE "PART B" and "PART C (Mandatory) Detailed Requirement" content**
- Extract content based on what's ACTUALLY across ALL processed sheets, regardless of domain
- Adapt language and focus to match the source content type
- If no content found for a subsection after checking ALL sheets, leave it blank
 
VALIDATION CHECK:
 
Before finalizing each section, verify that every piece of information can be traced back to the source requirements from the processed Excel sheets (excluding Test Scenarios). For Purpose and To-be process sections, ensure you've prioritized "PART B : (Mandatory) Detailed Requirement" content when available.
 
OUTPUT FORMAT:
Provide ONLY the markdown sections (## 1.0 Introduction, ### 1.1 Purpose, etc.) with the extracted content. Do not include any of these instructions, validation checks, or processing guidelines in your response.
 
""",
 
    "process_requirements": """
 
You are a Business Analyst expert creating sections 3.0–4.0 of a comprehensive BRD.
 
PREVIOUS CONTENT:
 
{previous_content}
 
SOURCE REQUIREMENTS:
 
{requirements}
 
EXCEL FILE PROCESSING INSTRUCTIONS:
 
**FOR EXCEL FILES (.xlsx/.xls):**
- Process ALL sheets EXCEPT "Test Scenarios" sheet
- Include data from sheets: "Requirement", "Ops Risk Assessment", and any other available sheets
- Extract workflow, process, and business rule information from ALL relevant columns and rows
- Look for step-by-step processes, business rules, and functional requirements across all sheets
 
CRITICAL INSTRUCTIONS:
 
- Extract information from ALL available sheets (except Test Scenarios sheet)
- Identify ACTUAL workflows, processes, and business rules from ANY relevant sheet
- Focus on step-by-step logic, conditions, and decision points mentioned across ALL processed sheets
- Adapt to any business domain (training, validation, integration, access control, etc.)
- Do NOT create, assume, or fabricate any content not explicitly present in the source
 
Create ONLY the following sections with detailed content in markdown:
 
## 3.0 Process / Data Flow diagram / Figma
 
Extract DETAILED workflow/process information from ALL processed sheets:
 
### 3.1 Workflow Description
 
Create step-by-step process based on what's described across ALL processed sheets:
- What triggers the process or workflow?
- What are the sequential steps or stages?
- What decision points, conditions, or validations occur?
- What are the different paths or outcomes?
- How are errors, exceptions, or edge cases handled?
- What user interactions or system responses are involved?
 
Format as logical flow:
- Step 1: [Action/Trigger from any relevant sheet]
  - If [condition mentioned in any sheet]: [result/next step]
  - If [alternative condition]: [alternative result]
- Step 2: [Next Action from any relevant sheet]
  - [Continue based on source content from processed sheets]
 
Look for process indicators across ALL sheets: "workflow", "process", "steps", "sequence", "flow", "journey", "condition", "if", "then", "when", "trigger", "action", "response"
 
## 4.0 Business / System Requirement
 
### 4.1 Functional Requirements
 
Module Name: [Extract exact application/module name from ANY processed sheet]
 
Create detailed requirement table based on content from ALL processed sheets:
 
| Rule ID| Rule Description | Expected Result| Dependency |
|-------------|---------------------|-------------------|----------------|
| 4.1.1 | [Extract specific business rule from ANY processed sheet] | [Exact expected behavior mentioned in ANY sheet] | [Technical/system dependencies noted in ANY sheet] |
 
Focus on extracting from ALL processed sheets:
- Specific business rules, validations, or logic mentioned
- Functional requirements and expected system behaviors
- User access controls, permissions, or restrictions
- Data validation, processing, or transformation rules
- Integration requirements and system interactions
 
### 4.2 System Requirements
 
Extract BUSINESS functional requirements from ALL processed sheets:
- Look for detailed requirement sections in the "Requirement" sheet primarily
- Also check other sheets for additional functional requirements
- Extract information from any columns containing requirement descriptions
- Include business rules, validation requirements, and functional specifications and it should releate more towards the technical side of things
 
**SPECIFIC FOR EXCEL:**
- If there's a "Requirement" sheet, prioritize extracting from detailed requirement sections
- Check for cells like "Detailed Requirement", "Business Rule", "Functional Spec", etc.
- Process other sheets for supplementary functional requirements
 
IMPORTANT:
 
- Use markdown headings
- Create detailed requirement tables with multiple columns
- Base all content on what's explicitly stated across ALL processed sheets
- Adapt terminology and focus to match the source domain
- Leave blank if no content found after checking ALL relevant sheets
 
VALIDATION CHECK:
 
Before finalizing each section, verify that every piece of information can be traced back to the source requirements from the processed Excel sheets (excluding Test Scenarios). Remove any content that cannot be directly attributed to the source documents.
 
OUTPUT FORMAT:
Provide ONLY the markdown sections (## 3.0, ### 3.1, etc.) with the extracted content. Do not include any of these instructions, validation checks, or processing guidelines in your response.
 
""",
 
    "data_communication": """
 
You are a Business Analyst expert creating section 5.0 of a comprehensive BRD.
 
PREVIOUS CONTENT:
 
{previous_content}
 
SOURCE REQUIREMENTS:
 
{requirements}
 
EXCEL FILE PROCESSING INSTRUCTIONS:
 
**FOR EXCEL FILES (.xlsx/.xls):**
- Process ALL sheets EXCEPT "Test Scenarios" sheet
- Include data from sheets: "Requirement", "Ops Risk Assessment", and any other available sheets
- Extract data requirements, specifications, and communication needs from ALL relevant sheets
- Look for data-related requirements across all processed sheets
 
CRITICAL INSTRUCTIONS:
 
- Extract information from ALL available sheets (except Test Scenarios sheet)
- Identify ACTUAL data and communication needs from ANY relevant sheet
- Adapt to any type of data requirements (user data, transaction data, training data, etc.)
- Do NOT create, assume, or fabricate any content not explicitly present in the source
 
Create ONLY the following section with detailed content in markdown:
 
## 5.0 MIS / DATA Requirement
 
### 5.1 Data Specifications
 
OUTPUT FORMAT (MANDATORY TABLE ONLY):
| Data Category | Specific Fields/Elements | Frequency/Trigger | Business Purpose |
|---------------|--------------------------|-------------------|------------------|
| [Extract from source] | [field1, field2, ...] | [e.g., Daily/On Event] | [purpose from source] |
 
RULES:
- OUTPUT ONLY the table above (no extra text)
** for filling table search for:**
- Data category,
- fields, or attributes needed
- frequency trigger like monthly , daily, weekkly
- Busines purpose : purpose of the data categories functionally and as per business logic.
 
### 5.2 Reporting and Analytics Needs
 
OUTPUT FORMAT (MANDATORY TABLE ONLY):
| Report/Dashboard Name | Visualization Type | Analytics Tool Suggestion | Target Audience | Frequency | Business Value |
|-----------------------|--------------------|---------------------------|-----------------|-----------|----------------|
| [Extract from source] | [e.g., Line chart] | [e.g., Tableau/Power BI]  | [e.g., Business users] | [e.g., Daily] | [value/goal from source] |
 
RULES:
- OUTPUT ONLY the table above (no extra text)
** for filling table search for:**
    -Reports, dashboards, or analytics required
    - Data visualization or presentation requirements
    - Best tools to build these plots/charts (specifically which BI tool)
    - User roles or audiences needing access
    - Frequency or scheduling of reports
    - business values.
 
 
### 5.3 Data Sources and Destinations
 
**IMPORTANT: Create a markdown table for data flow information found in the source requirements.**
 
Extract data flow information from ALL processed sheets and present in table format:
 
| Source System | Destination System | Data Type | Integration Method | Frequency | Dependencies |
|-------------------|------------------------|---------------|----------------------|---------------|------------------|
| [Extract from source] | [Extract from source] | [Extract from source] | [Extract from source] | [Extract from source] | [Extract from source] |
 
**Search for:**
- Source systems, databases, or applications providing data
- Target systems, repositories, or destinations for data
- Integration points, APIs, or data exchange mechanisms
- Data flow directions and transformation requirements
- External systems, third-party sources, or partner integrations
- Master data management or reference data needs
 
 
 
## 6.0 Communication Requirement
**VERY CRITICAL** : Never skip Communication requirement section.
 
**PRIORITY SEARCH STRATEGY:**
1. **FIRST PRIORITY**: Look specifically for "part_e" content and adjacent_content list in the priority_content section
2. **SECOND PRIORITY**: Search for "PART E : (Mandatory/Optional)" or similar patterns
 
**EXTRACTION INSTRUCTIONS:**
STEP BY STEP
   1.In part_e you will see list of content in that you will be having row no. and  text .
   2.from context list you have the questions in text :
      - Whether the any change has to be done in communication related to given modules"
      - IF YES, please specify the communication list
      -If YES, please confirm whether the communication format is attached in the call"
      - Please confirm whether necessary approvals taken on the communication format (HOD Approval, Legal approval)
      - To whom the communication has to be addressed
      - Mode of communication
  3.for all above text question you will be having row no.
  4. Now go to adjacent context lists, now match the same row no. of question to the row no. of adjacent context context list and text corresponding to that row in adjacet_context will be answer of that question.
 
**VERY CRITICAL**: Even if all the answers in the adjacent context list are no or blank, display the communications section in that case as well.
 
**OUTPUT Format **
   As you got all the questions and its respective answers in step 4.
   Now State the statements like , if your  answer for 1st question is no or blank then 1st statement will be -  No any changes has to be done in communication related to given modules.
   Now State the statements like , if your  answer 1st question is yes then 1st statement will be . The changes has to be done to be done in communication related to given modules.
   Now State the second statement , The list of communication is : if no list is given or answer for that row is blank then say no list is given in source document.
   similarly state other statement for all the questions.
   If your question;s answer is No that is no changes made in communication related modeule , then obviously there will be no list of communication so you skip the statemnt and go for rest of the statements.
   **CRITICAL**:- For line breaks, insert an actual Enter instead directly  giving literal '\n\n'
 
IMPORTANT:
 
- **MANDATORY: Create tables for sections 5.1, 5.2 and 5.3 using the specified formats above**
- Use markdown headings
- Extract content based on what's ACTUALLY across ALL processed sheets, regardless of domain
- Adapt language and focus to match the source content type
- If no content found for a subsection after checking ALL sheets, use the specified "not found" table format
- Preserve any existing tables in markdown format from ANY processed sheet
 
VALIDATION CHECK:
 
Before finalizing each section, verify that every piece of information can be traced back to the source requirements from the processed Excel sheets (excluding Test Scenarios). Remove any content that cannot be directly attributed to the source documents.
 
OUTPUT FORMAT:
Provide ONLY the markdown sections (## 5.0, ### 5.1, etc.) with the extracted content in TABLE FORMAT for sections 5.1, 5.2, and 5.3. Do not include any of these instructions, validation checks, or processing guidelines in your response.
 
""",
 
    "testing_final": """
 
You are a Business Analyst expert creating sections 7.0–11.0 of a comprehensive BRD.
 
PREVIOUS CONTENT:
 
{previous_content}
 
SOURCE REQUIREMENTS:
 
{requirements}
 
CRITICAL INSTRUCTIONS FOR ALL SECTIONS:
 
- Extract information ONLY from the provided source requirements
- For Test Scenarios: PRIORITY CHECK - First look for existing test scenarios in source
- Adapt to any business domain or requirement type
- Do NOT create, assume, or fabricate any content not explicitly present in the source
 
Create ONLY the following sections with detailed content in markdown:
 
## 7.0 Test Scenarios
 
**PRIMARY APPROACH - Extract Existing Test Scenarios:**
FIRST, thoroughly scan ALL source requirements documents for existing test content using these keywords:
- "Test Scenarios" / "Test Scenario"
- "Test Cases" / "Test Case"
- "Test case Scenarios"
- "Testing" / "Test Plan"
- "Verification" / "Validation"
 
**IF existing test scenarios/cases ARE FOUND in source documents:**
- Extract and preserve ALL the EXACT test scenarios from the source (require all the test scenarios from the source)
- Maintain original test structure, format, and content
- Convert to standardized markdown table format:
 
| Test ID | Test Scenario Name | Objective | Test Steps | Expected Results | Type |
|-------------|---------------|---------------|----------------|---------------------|----------|
| [Extract ID] | [Extract Name] | [Extract Objective] | [Extract Steps] | [Extract Results] | [Extract Type] |
 
Also, ADDING on to this, generate test scenarios based EXCLUSIVELY on functionality explicitly described in source requirements
 
**STOP HERE - Do not proceed to Secondary Approach if existing tests are found**
 
---
 
**SECONDARY APPROACH - Generate from Functional Requirements:**
**ONLY EXECUTE IF PRIMARY APPROACH YIELDS NO RESULTS**
 
IF NO existing test scenarios are found in ANY source documents, THEN generate test scenarios based EXCLUSIVELY on functionality explicitly described in source requirements:
 
| Test ID | Test Name | Objective | Test Steps | Expected Results | Type |
|-------------|---------------|---------------|----------------|---------------------|----------|
 
Create exactly 5 test scenarios covering:
- Primary functional requirements mentioned in source
- Different user roles, permissions, or access levels described  
- Various input conditions, data scenarios, or edge cases noted
- Error conditions, exceptions, or validation failures mentioned
- Integration points, API calls, or system interactions described
 
**CRITICAL:** Base ALL generated test scenarios ONLY on what is explicitly described in the source requirements. Do not infer or assume functionality not documented.
 
**EXECUTION RULE:** Use Primary Approach OR Secondary Approach - NEVER BOTH.
 
## 8.0 Questions / Suggestions
 
**SEARCH STRATEGY:**
- Scan ALL source documents for explicit questions, suggestions, or clarifications
- Look for keywords: "Question", "Clarification", "Unknown", "Assumption", "Dependency", "Suggestion", "Recommendation"
 
**IF questions/suggestions ARE FOUND in source:**
- List exact questions, clarifications, or unknowns from source
- List exact assumptions, dependencies, or prerequisites from source  
- List exact suggestions, recommendations, or enhancements from source
 
**IF NO questions/suggestions are found in source:**
- Leave this section completely BLANK
- Do NOT generate, create, or assume any questions or suggestions
- Do NOT infer potential issues or recommendations
 
**CRITICAL:** Only extract content that is explicitly stated in the source documents. Never generate, create, assume, or fabricate any questions, suggestions, or recommendations not present in the source.
 
## 9.0 Reference Document
 
List exact source documents
 
**CRITICAL:** Only extract content that is explicitly stated in the source documents. Never generate, create, assume, or fabricate anything
 
## 10.0 Appendix
 
**SEARCH STRATEGY:**
- Scan ALL source documents for explicit appendix and supporting information
- Look for keywords: "Appendix"
 
**IF appendix/supporting information ARE FOUND in source:**
- List exact appendix and supporting information
 
**IF NO appendix/supporting information are found in source:**
- Leave this section completely BLANK
- Do NOT generate, create, or assume any appendix or supporting information
- Do NOT infer potential supporting information
 
## 11.0 Risk Evaluation
 
**CRITICAL EXTRACTION RULE:**
Extract the EXACT table content from the source documents. Do NOT modify, interpret, or reformat the content.
 
**SEARCH FOR RISK CONTENT:**
- Look for any sheet named "Ops Risk Assessment", "Risk Evaluation", "Risk Assessment", or similar
- Look for any table or structured data containing risk-related information
- Search for keywords: "risk", "evaluation", "assessment", "impact", "controls"
 
**EXTRACTION PROCESS:**
1. **IF a risk table/data is found in the source:**
   - Copy the EXACT column headers from the source
   - Copy the EXACT row data from the source
   - Maintain the EXACT table structure and content
   - Convert to clean markdown table format WITHOUT changing any text content
   - Include ALL rows and columns as they appear in the source
 
2. **EXAMPLE - If source has this table:**
   ```
   | Risk Type | Impact | Mitigation | Status |
   | High Risk | Operational | Control A | Active |
   ```
 
   **OUTPUT EXACTLY:**
   ```
   | Risk Type | Impact | Mitigation | Status |
   |-----------|--------|------------|---------|
   | High Risk | Operational | Control A | Active |
   ```
 
3. **IF NO risk content found:**
   - State: "No Risk Evaluation content found in source documents"
   - Do NOT create any template or sample content
 
**FORBIDDEN:**
- Do NOT generate placeholder text like "List down the business risks"
- Do NOT create template structures
- Do NOT interpret or modify the source content
- Do NOT add explanatory text or instructions in table cells
 
**VALIDATION:**
Every piece of content in this section must be traceable to the exact source document content.
 
VALIDATION CHECK:
 
Before finalizing sections 8.0-11.0, verify that every piece of information can be traced back to the source requirements. Remove any content that cannot be directly attributed to the source documents.
 
OUTPUT FORMAT:
Provide ONLY the markdown sections (## 7.0, ### 7.1, etc.) with the extracted content. Do not include any of these instructions, validation checks, or processing guidelines in your response.
q
 
"""
 
}
 

def estimate_content_size(text):
    return len(text)

def chunk_requirements(requirements, max_chunk_size=8000):
    if estimate_content_size(requirements) <= max_chunk_size:
        return [requirements]
    
    sections = requirements.split('\n\n')
    chunks = []
    current_chunk = ""
    
    for section in sections:
        if estimate_content_size(current_chunk + section) > max_chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = section
        else:
            current_chunk += "\n\n" + section if current_chunk else section
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

@st.cache_resource
def initialize_sequential_chains(api_provider, api_key, azure_endpoint=None, azure_deployment=None, api_version=None):
    
    if api_provider == "OpenAI":
        model = ChatOpenAI(
            openai_api_key=api_key,
            model_name="gpt-3.5-turbo-16k",
            temperature=0.2,
            top_p=0.2
        )
    elif api_provider == "AzureOpenAI":
        model = AzureChatOpenAI(
            azure_endpoint=azure_endpoint,
            openai_api_key=api_key,
            azure_deployment=azure_deployment,
            api_version=api_version,
            temperature=0.2,
            top_p=0.2
        )
    else:
        model = ChatGroq(
            groq_api_key=api_key,
            model_name="llama3-70b-8192",
            temperature=0.2,
            top_p=0.2
        )
    
    chains = []
    chain1 = LLMChain(
        llm=model,
        prompt=PromptTemplate(
            input_variables=['requirements'],
            template=SECTION_TEMPLATES["intro_impact"]
        ),
        output_key="intro_impact_sections"
    )
    
    chain2 = LLMChain(
        llm=model,
        prompt=PromptTemplate(
            input_variables=['previous_content', 'requirements'],
            template=SECTION_TEMPLATES["process_requirements"]
        ),
        output_key="process_requirements_sections"
    )
    
    chain3 = LLMChain(
        llm=model,
        prompt=PromptTemplate(
            input_variables=['previous_content', 'requirements'],
            template=SECTION_TEMPLATES["data_communication"]
        ),
        output_key="data_communication_sections"
    )
    
    chain4 = LLMChain(
        llm=model,
        prompt=PromptTemplate(
            input_variables=['previous_content', 'requirements'],
            template=SECTION_TEMPLATES["testing_final"]
        ),
        output_key="testing_final_sections"
    )
    
    return [chain1, chain2, chain3, chain4]

def generate_brd_sequentially(chains, requirements):
    
    req_chunks = chunk_requirements(requirements)
    
    if len(req_chunks) > 1:
        st.info(f"Large content detected. Processing in {len(req_chunks)} chunks...")
    
    combined_requirements = "\n\n=== DOCUMENT BREAK ===\n\n".join(req_chunks)
    
    st.write("="*120)
    st.write("📋 COMBINED REQUIREMENTS SENT TO LLM:")
    st.write("="*120)

    product_alignment = load_product_alignment()
    # Load API catalog and append as reference block for the LLM (prompt-only usage)
    try:
        with open(os.path.join(os.path.dirname(__file__), "apis_full.json"), "r", encoding="utf-8") as _f:
            apis_catalog_json = json.load(_f)
    except Exception:
        apis_catalog_json = {}

    if product_alignment:
        product_alignment_text = "\n\n=== PRODUCT ALIGNMENT DATA ===\n"
        product_alignment_text += json.dumps(product_alignment, indent=2)
        product_alignment_text += "\n" + "="*50
        combined_requirements = combined_requirements + product_alignment_text
    
    if apis_catalog_json:
        combined_requirements += "\n\n=== KNOWN API CATALOG (READ-ONLY REFERENCE) ===\n"
        combined_requirements += json.dumps(apis_catalog_json, indent=2)
        combined_requirements += "\n" + "="*50
    
    with st.expander("📄 View Complete Requirements Content", expanded=False):
        st.text_area("Full Content", combined_requirements, height=400)
    
    st.write(f"📊 **Content Statistics:**")
    st.write(f"- Total characters: {len(combined_requirements):,}")
    lines_count = len(combined_requirements.split('\n'))
    words_count = len(combined_requirements.split())
    st.write(f"- Total lines: {lines_count:,}")
    st.write(f"- Total words (approx): {words_count:,}")
    st.write(f"- Number of chunks: {len(req_chunks)}")
    
    st.write(f"📖 **Content Preview (First 2000 characters):**")
    st.code(combined_requirements[:2000] + "..." if len(combined_requirements) > 2000 else combined_requirements)
    
    sections = [line for line in combined_requirements.split('\n') if line.strip().startswith('===')]
    if sections:
        st.write(f"**Document Structure:**")
        for section in sections[:10]:
            st.write(f"- {section.strip()}")
        if len(sections) > 10:
            st.write(f"- ... and {len(sections) - 10} more sections")
    
    st.write("="*120)
    
    previous_content = ""
    final_sections = []
    
    for i, chain in enumerate(chains):
        try:
            st.write(f"\\n🔗 **PROCESSING CHAIN {i+1}/4**")
            st.write(f"{'='*60}")
            
            with st.expander(f"🔍 Chain {i+1} Details - Click to expand", expanded=False):
                
                if i == 0:
                    st.write("**Input to Chain 1 (Introduction & Impact Analysis):**")
                    st.write(f"- Requirements length: {len(combined_requirements):,} characters")
                    st.write("**Requirements Preview:**")
                    st.code(combined_requirements[:1000] + "..." if len(combined_requirements) > 1000 else combined_requirements)
                    
                    st.write("**Template Used:**")
                    st.code(SECTION_TEMPLATES["intro_impact"][:500] + "...")
                    
                    result = chain.run(requirements=combined_requirements)
                else:
                    chain_names = ["", "Process & Requirements", "Data & Communication", "Testing & Final"]
                    st.write(f"**Input to Chain {i+1} ({chain_names[i]}):**")
                    st.write(f"- Previous content length: {len(previous_content):,} characters")
                    st.write(f"- Requirements length: {len(combined_requirements):,} characters")
                    
                    st.write("**Previous Content Preview:**")
                    st.code(previous_content[:800] + "..." if len(previous_content) > 800 else previous_content)
                    
                    st.write("**Requirements Preview:**")
                    st.code(combined_requirements[:800] + "..." if len(combined_requirements) > 800 else combined_requirements)
                    
                    template_keys = ["", "process_requirements", "data_communication", "testing_final"]
                    st.write(f"**Template Used ({template_keys[i]}):**")
                    st.code(SECTION_TEMPLATES[template_keys[i]][:500] + "...")
                    
                    result = chain.run(previous_content=previous_content, requirements=combined_requirements)
                
                st.write(f"**Chain {i+1} Output:**")
                st.write(f"- Response length: {len(result):,} characters")
                result_lines = len(result.split('\n'))
                result_words = len(result.split())
                st.write(f"- Response lines: {result_lines:,}")
                st.write(f"- Response words (approx): {result_words:,}")
                
                output_sections = [line for line in result.split('\n') if line.strip().startswith('##')]
                if output_sections:
                    st.write("**Sections Generated:**")
                    for section in output_sections:
                        st.write(f"- {section.strip()}")
                
                st.write("**Response Preview:**")
                st.code(result[:1000] + "..." if len(result) > 1000 else result)
            
            print(f"\n{'='*60}")
            print(f"CHAIN {i+1} INPUT:")
            print(f"{'='*60}")
            
            if i == 0:
                print("Input to Chain 1 (intro_impact):")
                print(f"Requirements length: {len(combined_requirements)} characters")
                print("First 1000 characters of requirements:")
                print(combined_requirements[:1000] + "..." if len(combined_requirements) > 1000 else combined_requirements)
                
            else:
                print(f"Input to Chain {i+1}:")
                print(f"Previous content length: {len(previous_content)} characters")
                print(f"Requirements length: {len(combined_requirements)} characters")
                print("Previous content (first 500 chars):")
                print(previous_content[:500] + "..." if len(previous_content) > 500 else previous_content)
                print("\nRequirements (first 500 chars):")
                print(combined_requirements[:500] + "..." if len(combined_requirements) > 500 else combined_requirements)

            if i == 0 and product_alignment:
                result = expand_product_categories(result, product_alignment)
            
            # Removed API injection: rely on prompt with catalog JSON only
            
            print(f"\nCHAIN {i+1} OUTPUT:")
            print(f"Response length: {len(result)} characters")
            print("First 1000 characters of response:")
            print(result[:1000] + "..." if len(result) > 1000 else result)
            print(f"{'='*60}")
            
            final_sections.append(result)
            previous_content += "\\n\\n" + result
            
            st.write(f"✅ **Completed section group {i+1}/4**")
            st.write(f"📈 **Cumulative content length: {len(previous_content):,} characters**")
            
        except Exception as e:
            print(f"ERROR in chain {i+1}: {str(e)}")
            st.error(f"❌ Error in chain {i+1}: {str(e)}")
            final_sections.append(f"## Error in section group {i+1}\\nError processing this section: {str(e)}")
    
    final_brd = "\\n\\n".join(final_sections)
    
    st.write("\\n" + "="*80)
    st.write("📋 **FINAL BRD GENERATION COMPLETE**")
    st.write("="*80)
    
    with st.expander("📊 Final BRD Statistics & Preview", expanded=True):
        st.write(f"**Final Statistics:**")
        st.write(f"- Total final BRD length: {len(final_brd):,} characters")
        final_lines = len(final_brd.split('\n'))
        final_words = len(final_brd.split())
        st.write(f"- Total lines: {final_lines:,}")
        st.write(f"- Total words (approx): {final_words:,}")
        
        final_sections_headers = [line for line in final_brd.split('\n') if line.strip().startswith('##')]
        if final_sections_headers:
            st.write(f"**Generated Sections ({len(final_sections_headers)}):**")
            for section in final_sections_headers:
                st.write(f"- {section.strip()}")
        
        st.write("**Final BRD Preview (first 2000 characters):**")
        st.code(final_brd[:2000] + "..." if len(final_brd) > 2000 else final_brd)
    
    print("\n" + "="*80)
    print("FINAL BRD CONTENT:")
    print("="*80)
    print(f"Total final BRD length: {len(final_brd)} characters")
    print("Final BRD (first 2000 characters):")
    print(final_brd[:2000] + "..." if len(final_brd) > 2000 else final_brd)
    print("="*80)
    
    return final_brd

def create_toc_styles(doc):
    styles = doc.styles
    
    try:
        toc1_style = styles['TOC 1']
    except KeyError:
        toc1_style = styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
        toc1_style.font.name = 'Calibri'
        toc1_style.font.size = Pt(11)
        toc1_style.paragraph_format.left_indent = Inches(0)
        toc1_style.paragraph_format.space_after = Pt(0)
    
    try:
        toc2_style = styles['TOC 2']
    except KeyError:
        toc2_style = styles.add_style('TOC 2', WD_STYLE_TYPE.PARAGRAPH)
        toc2_style.font.name = 'Calibri'
        toc2_style.font.size = Pt(11)
        toc2_style.paragraph_format.left_indent = Inches(0.25)
        toc2_style.paragraph_format.space_after = Pt(0)

def create_clickable_toc(doc):
    toc_heading = doc.add_heading('Table of Contents', level=1)
    add_bookmark(toc_heading, 'TOC')
    
    create_toc_styles(doc)
    
    toc_entries = [
        ("1.0 Introduction", "introduction"),
        ("    1.1 Purpose", "purpose"),
        ("    1.2 As-is process", "process_solution"),
        ("    1.3 To be process / High level solution", "process_solution"),
        ("2.0 Impact Analysis", "impact_analysis"),
        ("    2.1 Impacted Products", "impacted_products"),
        ("    2.2 Applications Impacted", "applications_impacted"), 
        ("    2.3 List of APIs required", "apis_required"),
        ("3.0 Process / Data Flow diagram / Figma", "process_flow"),
        ("4.0 Business / System Requirement", "business_requirements"),
        ("5.0 MIS / DATA Requirement", "mis_data_requirement"),
        ("6.0 Communication Requirement", "communication_requirement"),
        ("7.0 Test Scenarios", "test_scenarios"),
        ("8.0 Questions / Suggestions", "questions_suggestions"),
        ("9.0 Reference Document", "reference_document"),
        ("10.0 Appendix", "appendix"),
        ("11.0 Risk Evaluation", "risk_evaluation")
    ]

    bookmark_mapping = {}
    for entry_text, bookmark_name in toc_entries:
        bookmark_mapping[bookmark_name] = entry_text
        
    for entry_text, bookmark_name in toc_entries:
        toc_paragraph = doc.add_paragraph()
        
        try:
            if entry_text.startswith("    "):
                toc_paragraph.style = 'TOC 2'
            else:
                toc_paragraph.style = 'TOC 1'
        except KeyError:
            if entry_text.startswith("    "):
                toc_paragraph.paragraph_format.left_indent = Inches(0.25)
            toc_paragraph.paragraph_format.space_after = Pt(0)
        
        if entry_text.startswith("    "):
            toc_paragraph.add_run("    ")
            link_text = entry_text.strip()
        else:
            link_text = entry_text
            
        add_hyperlink(toc_paragraph, link_text, bookmark_name, is_internal=True)
        
        toc_paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
        
        toc_paragraph.add_run("\t")
        
        page_run = toc_paragraph.add_run()
        
        fldChar_begin = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>')
        instrText = parse_xml(f'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGEREF {bookmark_name} \\h </w:instrText>')
        fldChar_end = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>')
        
        page_run._r.append(fldChar_begin)
        page_run._r.append(instrText)
        page_run._r.append(fldChar_end)
        
        page_run.add_text(" ")
    
    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_run = note_para.add_run("IMPORTANT: ")
    note_run.bold = True
    note_run.font.color.rgb = RGBColor(255, 0, 0)
    
    note_para.add_run("To see actual page numbers in this Table of Contents:")
    note_para.add_run("Press 'Ctrl + A' to select all, then F9 to update all fields in the document.")
    
    return bookmark_mapping

def add_hyperlink(paragraph, text, url_or_bookmark, is_internal=True):
    part = paragraph.part
    
    hyperlink = OxmlElement('w:hyperlink')
    
    if is_internal:
        hyperlink.set(qn('w:anchor'), url_or_bookmark)
    else:
        r_id = part.relate_to(url_or_bookmark, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    new_run.append(rPr)
    
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    
    return hyperlink

def add_bookmark(paragraph, bookmark_name):
    bookmark_id = str(abs(hash(bookmark_name)) % 1000000)
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)

def add_section_with_bookmark(doc, heading_text, bookmark_name, level=1):
    heading = doc.add_heading(heading_text, level=level)
    add_bookmark(heading, bookmark_name)
    
    return heading



def parse_markdown_table(table_text):
    def clean_cell_value(cell_text):
        if cell_text is None:
            return ""
        
        str_val = str(cell_text).strip()
        
        if str_val.lower() == 'nan':
            return ""
        
        if str_val.startswith("Unnamed"):
            return "Insert Column Name"
        
        return str_val
    
    lines = [line.strip() for line in table_text.split('\n') if line.strip()]
    
    if len(lines) < 2:
        return None
    
    # Remove separator line (the one with ---)
    filtered_lines = []
    for line in lines:
        if not re.match(r'^[\s\|\-]+$', line):  # Skip lines that only contain |, -, and spaces
            filtered_lines.append(line)
    
    if len(filtered_lines) < 2:  # Need at least header + 1 data row
        return None
    
    table_data = []
    
    for line in filtered_lines:
        # Clean up the line - remove leading/trailing pipes
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        
        # Split by | and clean each cell
        cells = [clean_cell_value(cell.strip()) for cell in line.split('|')]
        
        # Remove empty cells from the end
        while cells and not cells[-1]:
            cells.pop()
        
        if cells:  # Only add if there are non-empty cells
            table_data.append(cells)
    
    if not table_data:
        return None
    
    # Ensure all rows have the same number of columns as the header
    if len(table_data) > 0:
        max_cols = len(table_data[0])  # Use header row as reference
        
        # Normalize all rows to have the same number of columns
        normalized_data = []
        for i, row in enumerate(table_data):
            if i == 0:  # Header row
                normalized_row = row[:max_cols]  # Don't extend header, just truncate if needed
            else:  # Data rows
                # Extend with empty strings if needed, truncate if too long
                if len(row) < max_cols:
                    normalized_row = row + [''] * (max_cols - len(row))
                else:
                    normalized_row = row[:max_cols]
            normalized_data.append(normalized_row)
        
        return normalized_data
    
    return None

def create_table_in_doc(doc, table_data):
    def clean_table_cell_value(cell_text):
        if cell_text is None:
            return ""
        
        str_val = str(cell_text).strip()
        
        if str_val.lower() == 'nan':
            return ""
        
        if str_val.startswith("Unnamed"):
            return "Insert Column Name"
        
        return str_val
    
    if not table_data or len(table_data) < 1:
        return None
    
    # Filter out completely empty columns
    filtered_data = []
    if len(table_data) > 0:
        num_cols = len(table_data[0])
        
        # Check which columns have actual data
        columns_with_data = []
        for col_idx in range(num_cols):
            has_data = False
            for row in table_data:
                if col_idx < len(row) and clean_table_cell_value(row[col_idx]):
                    has_data = True
                    break
            if has_data:
                columns_with_data.append(col_idx)
        
        # Create filtered table data with only columns that have data
        for row in table_data:
            filtered_row = []
            for col_idx in columns_with_data:
                if col_idx < len(row):
                    filtered_row.append(clean_table_cell_value(row[col_idx]))
                else:
                    filtered_row.append("")
            if filtered_row:  # Only add if row has content
                filtered_data.append(filtered_row)
    
    if not filtered_data or len(filtered_data) < 1:
        return None
    
    table = doc.add_table(rows=len(filtered_data), cols=len(filtered_data[0]))
    table.style = 'Table Grid'
    
    # Style header row
    for i, cell_text in enumerate(filtered_data[0]):
        if i < len(table.rows[0].cells):
            cell = table.rows[0].cells[i]
            cell.text = cell_text if cell_text else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Fill data rows
    for row_idx, row_data in enumerate(filtered_data[1:], 1):
        if row_idx < len(table.rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_text if cell_text else ""
    
    return table

def extract_content_from_docx(doc_file):
    doc = Document(doc_file)
    content = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text.strip())
    
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_content.append(" | ".join(row_text))
        if table_content:
            content.append("TABLE:")
            content.append("\n".join(table_content))
    
    return "\n".join(content)

def extract_content_from_pdf(pdf_file):
    content = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                content.append(page.extract_text())
            
            tables = page.extract_tables()
            for table in tables:
                if table:
                    table_text = []
                    for row in table:
                        row_text = [str(cell) if cell else "" for cell in row]
                        table_text.append(" | ".join(row_text))
                    if table_text:
                        content.append("TABLE:")
                        content.append("\n".join(table_text))
    
    return "\n".join(content)

def extract_content_from_excel(excel_file, max_rows_per_sheet=70, max_sample_rows=10, visible_only=True):
    def clean_cell_value(cell_text):
        if cell_text is None:
            return "-"
        
        str_val = str(cell_text).strip()
        
        if str_val.lower() == 'nan':
            return "-"
        
        if str_val.startswith("Unnamed"):
            return "Insert Column Name"
        
        return str_val
    
    def extract_horizontal_table(df, trigger_row_idx, start_col_idx, table_identifier):
        """Extract horizontal table structure starting from a trigger row"""
        table_data = {
            "table_type": table_identifier,
            "headers": [],
            "data_rows": [],
            "raw_structure": []
        }
        
        try:
            # Look for header row - it should be the next non-empty row after trigger
            header_row_idx = None
            data_start_row_idx = None
            
            # Search for the actual header row (next 3 rows max)
            for search_offset in range(1, 4):
                check_row_idx = trigger_row_idx + search_offset
                if check_row_idx >= len(df):
                    break
                
                # Get the row data starting from the same column as trigger
                row_data = df.iloc[check_row_idx, start_col_idx:].values
                non_empty_count = sum(1 for cell in row_data if pd.notna(cell) and str(cell).strip())
                
                # If we find a row with multiple non-empty cells, it's likely the header
                if non_empty_count >= 3:  # At least 3 columns for a proper table
                    header_row_idx = check_row_idx
                    data_start_row_idx = check_row_idx + 1
                    break
            
            if header_row_idx is None:
                return table_data
            
            # Extract headers from the identified header row
            header_row = df.iloc[header_row_idx]
            headers = []
            header_col_indices = []
            
            # Start from the trigger column and move right to collect headers
            for col_idx in range(start_col_idx, len(header_row)):
                cell_value = header_row.iloc[col_idx]
                if pd.notna(cell_value) and str(cell_value).strip():
                    clean_header = clean_cell_value(cell_value)
                    if clean_header not in ["-", ""]:
                        headers.append(clean_header)
                        header_col_indices.append(col_idx)
                elif headers:  # Stop when we hit empty cells after finding headers
                    break
            
            table_data["headers"] = headers
            
            if not headers:
                return table_data
            
            # Extract data rows (typically just one row for this type of table)
            max_data_rows = min(data_start_row_idx + 3, len(df))  # Check next 3 rows max
            
            for data_row_idx in range(data_start_row_idx, max_data_rows):
                if data_row_idx >= len(df):
                    break
                
                data_row = df.iloc[data_row_idx]
                row_values = []
                has_meaningful_data = False
                
                # Extract values for each header column
                for col_idx in header_col_indices:
                    if col_idx < len(data_row):
                        cell_val = clean_cell_value(data_row.iloc[col_idx])
                        row_values.append(cell_val)
                        # Check for meaningful data (not just "-" or empty)
                        if cell_val not in ["-", "", "nan"]:
                            has_meaningful_data = True
                    else:
                        row_values.append("-")
                
                if has_meaningful_data:
                    # Use the first column value as row description, or create a generic one
                    row_description = "Data Row"
                    if start_col_idx > 0:
                        desc_cell = data_row.iloc[start_col_idx - 1] if start_col_idx - 1 < len(data_row) else None
                        if pd.notna(desc_cell):
                            row_description = clean_cell_value(desc_cell)
                    
                    row_data = {
                        "row_description": row_description,
                        "values": dict(zip(headers, row_values))
                    }
                    table_data["data_rows"].append(row_data)
            
            # Create markdown table if we have data
            if headers and table_data["data_rows"]:
                table_data["raw_structure"] = {
                    "markdown_table": create_markdown_table(headers, table_data["data_rows"]),
                    "structured_data": table_data["data_rows"]
                }
            
        except Exception as e:
            table_data["error"] = str(e)
        
        return table_data
    
    def create_markdown_table(headers, data_rows):
        if not headers or not data_rows:
            return ""
        
        header_line = "| " + " | ".join(headers) + " |"
        separator_line = "|" + "|".join([" --- " for _ in headers]) + "|"
        
        data_lines = []
        for row in data_rows:
            values = [row["values"].get(header, "-") for header in headers]
            data_line = "| " + " | ".join(values) + " |"
            data_lines.append(data_line)
        
        return "\n".join([header_line, separator_line] + data_lines)
    
    def find_and_extract_part_c_tables(df, part_c_row_idx, part_c_col_idx):
        """Find and extract both Products Impacted and Applications Impacted tables"""
        tables = []
        
        # Search in a wider range after the Part C header
        search_range = min(part_c_row_idx + 25, len(df))
        
        for search_row in range(part_c_row_idx + 1, search_range):
            if search_row >= len(df):
                break
            
            # Check the entire row for table triggers
            for col_idx in range(len(df.columns)):
                search_cell = df.iloc[search_row, col_idx]
                if pd.notna(search_cell):
                    cell_text = str(search_cell).strip()
                    
                    # Look for "Products Impacted" trigger
                    if "Products Impacted" in cell_text:
                        products_table = extract_horizontal_table(df, search_row, col_idx, "Products Impacted")
                        if products_table["headers"]:
                            tables.append(products_table)
                    
                    # Look for "Applications Impacted" trigger  
                    elif "Applications Impacted" in cell_text:
                        apps_table = extract_horizontal_table(df, search_row, col_idx, "Applications Impacted")
                        if apps_table["headers"]:
                            tables.append(apps_table)
        
        return tables
    
    result = {
        "metadata": {
            "total_sheets": 0,
            "processing_status": "success",
            "visible_only": visible_only,
            "max_rows_per_sheet": max_rows_per_sheet,
            "max_sample_rows": max_sample_rows
        },
        "priority_content": {
            "part_b": [],
            "part_c": [],
            "part_e": [] 
        },
        "sheets": [],
        "summary": {
            "part_b_found": False,
            "part_c_found": False,
            "part_e_found": False,
            "total_rows_processed": 0,
            "total_columns_processed": 0,
            "detailed_requirements_found": False
        }
    }
    
    try:
        if visible_only:
            wb = load_workbook(excel_file)
            visible_sheets = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                if sheet.sheet_state == 'visible':
                    visible_sheets.append(sheet_name)
            
            if visible_sheets:
                excel_data = pd.read_excel(excel_file, sheet_name=visible_sheets)
            else:
                result["metadata"]["processing_status"] = "error"
                result["metadata"]["error"] = "No visible sheets found in the Excel file"
                return json.dumps(result, indent=2)
            
            if not isinstance(excel_data, dict):
                excel_data = {visible_sheets[0]: excel_data}
        else:
            excel_data = pd.read_excel(excel_file, sheet_name=None)
        
        result["metadata"]["total_sheets"] = len(excel_data)
        
        # Extract Part C content with improved table detection
        for sheet_name, df in excel_data.items():
            if df.empty:
                continue
                
            for col_idx, col in enumerate(df.columns):
                for row_idx, cell_value in enumerate(df[col]):
                    cell_str = str(cell_value).strip()
                    part_c_patterns = [
                        "PART C : (Mandatory) Detailed Requirement",
                        "PART C (Mandatory) Detailed Requirement", 
                        "PART C : Mandatory Detailed Requirement",
                        "PART C Mandatory Detailed Requirement",
                        "PART C : Detailed Requirement",
                        "PART C Detailed Requirement",
                        "PART C:",
                        "Part C :",
                        "Part C:",
                        "PART C"
                    ]
                    
                    for pattern in part_c_patterns:
                        if pattern.lower() in cell_str.lower():
                            part_c_entry = {
                                "sheet_name": sheet_name,
                                "column": col,
                                "row": row_idx + 2,
                                "header": cell_str,
                                "content": [],
                                "adjacent_content": [],
                                "horizontal_tables": []
                            }
                            
                            # Extract content below Part C header
                            for next_row in range(row_idx + 1, min(row_idx + 10, len(df))):
                                if next_row < len(df):
                                    next_cell = df.iloc[next_row][col]
                                    if pd.notna(next_cell) and str(next_cell).strip():
                                        part_c_entry["content"].append({
                                            "row": next_row + 2,
                                            "text": str(next_cell).strip()
                                        })
                            
                            # Extract adjacent content
                            for adj_col_offset in [-1, 1]:
                                adj_col_index = col_idx + adj_col_offset
                                if 0 <= adj_col_index < len(df.columns):
                                    adj_col = df.columns[adj_col_index]
                                    for adj_row in range(max(0, row_idx-2), min(row_idx + 8, len(df))):
                                        adj_cell = df.iloc[adj_row][adj_col]
                                        if pd.notna(adj_cell) and str(adj_cell).strip():
                                            part_c_entry["adjacent_content"].append({
                                                "column": adj_col,
                                                "row": adj_row + 2,
                                                "text": str(adj_cell).strip()
                                            })
                            
                            # Use the improved table extraction function
                            tables = find_and_extract_part_c_tables(df, row_idx, col_idx)
                            part_c_entry["horizontal_tables"] = tables
                            
                            result["priority_content"]["part_c"].append(part_c_entry)
                            result["summary"]["part_c_found"] = True
                            break
        
        # Extract Part B content (unchanged)
        for sheet_name, df in excel_data.items():
            if df.empty:
                continue
                
            for col_idx, col in enumerate(df.columns):
                for row_idx, cell_value in enumerate(df[col]):
                    cell_str = str(cell_value).strip()
                    part_b_patterns = [
                        "PART B : (Mandatory) Detailed Requirement",
                        "PART B (Mandatory) Detailed Requirement", 
                        "PART B : Mandatory Detailed Requirement",
                        "PART B Mandatory Detailed Requirement",
                        "PART B : Detailed Requirement",
                        "PART B Detailed Requirement",
                        "PART B:",
                        "Part B :",
                        "Part B:",
                        "PART B"
                    ]
                    
                    for pattern in part_b_patterns:
                        if pattern.lower() in cell_str.lower():
                            part_b_entry = {
                                "sheet_name": sheet_name,
                                "column": col,
                                "row": row_idx + 2,
                                "header": cell_str,
                                "content": [],
                                "adjacent_content": []
                            }
                            
                            for next_row in range(row_idx + 1, min(row_idx + 10, len(df))):
                                if next_row < len(df):
                                    next_cell = df.iloc[next_row][col]
                                    if pd.notna(next_cell) and str(next_cell).strip():
                                        part_b_entry["content"].append({
                                            "row": next_row + 2,
                                            "text": str(next_cell).strip()
                                        })
                            
                            for adj_col_offset in [-1, 1]:
                                adj_col_index = col_idx + adj_col_offset
                                if 0 <= adj_col_index < len(df.columns):
                                    adj_col = df.columns[adj_col_index]
                                    for adj_row in range(max(0, row_idx-2), min(row_idx + 8, len(df))):
                                        adj_cell = df.iloc[adj_row][adj_col]
                                        if pd.notna(adj_cell) and str(adj_cell).strip():
                                            part_b_entry["adjacent_content"].append({
                                                "column": adj_col,
                                                "row": adj_row + 2,
                                                "text": str(adj_cell).strip()
                                            })
                            
                            result["priority_content"]["part_b"].append(part_b_entry)
                            result["summary"]["part_b_found"] = True
                            break

        # Extract Part E content (unchanged)
        part_e_patterns = [
            "PART E : (Mandatory/Optional)",
            "PART E (Mandatory/Optional)",
            "PART E:",
            "Part E :",
            "Part E:",
            "PART E"
        ]

        for sheet_name, df in excel_data.items():
            if df.empty:
                continue

            for col_idx, col in enumerate(df.columns):
                for row_idx, cell_value in enumerate(df[col]):
                    cell_str = str(cell_value).strip()

                    for pattern in part_e_patterns:
                        if pattern.lower() in cell_str.lower():
                            part_e_entry = {
                                "sheet_name": sheet_name,
                                "column": col,
                                "row": row_idx + 2,
                                "header": cell_str,
                                "content": [],
                                "adjacent_content": [],
                                "detailed_responses": []
                            }

                            # Collect exactly next 8 rows
                            for offset in range(1, 9):
                                next_row = row_idx + offset
                                if next_row < len(df):
                                    next_cell = df.iloc[next_row][col]
                                    part_e_entry["content"].append({
                                        "row": next_row + 2,
                                        "text": "" if pd.isna(next_cell) else str(next_cell).strip()
                                    })

                            # Collect adjacent values for same 8 rows
                            for adj_col_offset in [-1, 1]:
                                adj_col_index = col_idx + adj_col_offset
                                if 0 <= adj_col_index < len(df.columns):
                                    adj_col = df.columns[adj_col_index]
                                    for offset in range(1, 9):
                                        next_row = row_idx + offset
                                        if next_row < len(df):
                                            adj_cell = df.iloc[next_row][adj_col]
                                            part_e_entry["adjacent_content"].append({
                                                "column": adj_col,
                                                "row": next_row + 2,
                                                "text": "" if pd.isna(adj_cell) else str(adj_cell).strip()
                                            })

                            result["priority_content"]["part_e"].append(part_e_entry)
                            result["summary"]["part_e_found"] = True
                            break

        # Process remaining sheets (unchanged)
        for sheet_name, df in excel_data.items():
            if df.empty:
                continue
            
            original_row_count = len(df)
            if max_rows_per_sheet and len(df) > max_rows_per_sheet:
                df = df.head(max_rows_per_sheet)
            
            sheet_data = {
                "sheet_name": sheet_name,
                "dimensions": {
                    "rows": original_row_count,
                    "columns": len(df.columns),
                    "processed_rows": len(df)
                },
                "columns": {
                    "names": [clean_cell_value(col) for col in df.columns.tolist()],
                    "data_types": {clean_cell_value(col): str(dtype) for col, dtype in df.dtypes.to_dict().items()},
                    "numeric_columns": [clean_cell_value(col) for col in df.select_dtypes(include=['number']).columns.tolist()],
                    "key_columns": []
                },
                "sample_data": [],
                "detailed_requirements": [],
                "data_summary": {
                    "missing_data": {},
                    "unique_value_counts": {}
                }
            }
            
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['id', 'name', 'title', 'status', 'type', 'category', 'priority', 'requirement']):
                    sheet_data["columns"]["key_columns"].append(clean_cell_value(col))
            
            sample_size = min(max_sample_rows, len(df))
            if sample_size > 0:
                display_df = df.head(sample_size)
                
                for _, row in display_df.iterrows():
                    row_data = {}
                    for col, val in row.items():
                        cleaned_val = clean_cell_value(val)
                        if len(cleaned_val) > 50:
                            cleaned_val = cleaned_val[:47] + "..."
                        row_data[clean_cell_value(col)] = cleaned_val
                    sheet_data["sample_data"].append(row_data)
            
            for col in df.columns:
                col_str = str(col).lower()
                if any(keyword in col_str for keyword in ['requirement', 'detailed', 'description', 'specification']):
                    req_column = {
                        "column_name": clean_cell_value(col),
                        "requirements": []
                    }
                    
                    for idx, cell_value in enumerate(df[col]):
                        if pd.notna(cell_value) and str(cell_value).strip():
                            cell_text = str(cell_value).strip()
                            if len(cell_text) > 10:
                                req_column["requirements"].append({
                                    "row": idx + 2,
                                    "text": cell_text
                                })
                    
                    if req_column["requirements"]:
                        sheet_data["detailed_requirements"].append(req_column)
                        result["summary"]["detailed_requirements_found"] = True
            
            for col in sheet_data["columns"]["key_columns"][:3]:
                if col in df.columns and df[col].dtype == 'object':
                    unique_vals = df[col].dropna().unique()
                    if len(unique_vals) <= 20:
                        sheet_data["data_summary"]["unique_value_counts"][col] = [clean_cell_value(val) for val in unique_vals[:10]]
                    else:
                        sheet_data["data_summary"]["unique_value_counts"][col] = f"{len(unique_vals)} unique values"
            
            missing_data = df.isnull().sum()
            if missing_data.sum() > 0:
                missing_cols = missing_data[missing_data > 0].head(5)
                sheet_data["data_summary"]["missing_data"] = {clean_cell_value(col): int(count) for col, count in missing_cols.items()}
            
            result["sheets"].append(sheet_data)
            result["summary"]["total_rows_processed"] += len(df)
            result["summary"]["total_columns_processed"] += len(df.columns)
    
    except Exception as e:
        result["metadata"]["processing_status"] = "error"
        result["metadata"]["error"] = str(e)
    
    return json.dumps(result, indent=2, ensure_ascii=False)

def extract_content_from_msg(msg_file):
    try:
        temp_file = BytesIO(msg_file.getvalue())
        temp_file.name = msg_file.name
        
        msg = extract_msg.Message(temp_file)
        body_content = msg.body
        
        cleaned_body = re.sub(r'^(From|To|Cc|Subject|Sent|Date):.*?\n', '', body_content, flags=re.MULTILINE)
        cleaned_body = re.sub(r'_{10,}[\s\S]*$', '', cleaned_body)
        cleaned_body = re.sub(r'-{10,}[\s\S]*$', '', cleaned_body)
        
        disclaimer_pattern = r'DISCLAIMER:[\s\S]*?customercare@bajajallianz\.co\.in'
        cleaned_body = re.sub(disclaimer_pattern, '', cleaned_body)
        
        return cleaned_body.strip()
    
    except Exception as e:
        st.error(f"Error processing MSG file: {str(e)}")
        return ""

def add_header_with_logo(doc, logo_bytes):
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    
    run = header_para.add_run()
    logo_stream = BytesIO(logo_bytes)
    run.add_picture(logo_stream, width=Inches(1.5))

def create_word_document(content, logo_data=None):
    doc = Document()
    
    if logo_data:
        add_header_with_logo(doc, logo_data)
    
    for _ in range(12):
        doc.add_paragraph()
    
    title = doc.add_heading('Business Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    doc.add_heading('Version History', level=1)
    version_table = doc.add_table(rows=5, cols=5)
    version_table.style = 'Table Grid'
    hdr_cells = version_table.rows[0].cells
    headers = ['Version', 'Date', 'Author', 'Change description', 'Review by']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    doc.add_paragraph('**To be reviewed and filled in by IT Team.**')
    
    doc.add_heading('Sign-off Matrix', level=1)
    signoff_table = doc.add_table(rows=5, cols=5)
    signoff_table.style = 'Table Grid'
    hdr_cells = signoff_table.rows[0].cells
    headers = ['Version', 'Sign-off Authority', 'Business Function', 'Sign-off Date', 'Email Confirmation']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    doc.add_page_break()
    
    bookmark_mapping = create_clickable_toc(doc)
    if bookmark_mapping is None:
        bookmark_mapping = {}
    
    doc.add_page_break()
    
    sections = content.split('##')
    
    introduction_started = False

    for i, section in enumerate(sections):
        if section.strip():
            lines = section.strip().split('\n')
            if lines:
                heading_line = lines[0].strip()
            
                bookmark_name = None
                for bookmark, heading_text in bookmark_mapping.items():
                    if heading_line.lower().replace('#', '').strip() in heading_text.lower():
                        bookmark_name = bookmark
                        break
            
                if heading_line.startswith('###'):
                    level = 2
                    heading_text = heading_line.replace('###', '').strip()
                else:
                    level = 1
                    heading_text = heading_line.replace('##', '').strip()
            
                section_name_lower = heading_text.lower()
                if 'introduction' in section_name_lower or section_name_lower.startswith('1.0'):
                    introduction_started = True
            
                if level == 1 and i > 0 and not introduction_started:
                    doc.add_page_break()
            
                if bookmark_name:
                    add_section_with_bookmark(doc, heading_text, bookmark_name, level)
                else:
                    doc.add_heading(heading_text, level)
            
                j = 1
                while j < len(lines):
                    line = lines[j].strip()
                
                    if line and '|' in line and line.count('|') >= 2:
                        table_lines = []
                        while j < len(lines) and lines[j].strip() and '|' in lines[j]:
                            table_lines.append(lines[j].strip())
                            j += 1
                    
                        if table_lines:
                            table_data = parse_markdown_table('\n'.join(table_lines))
                            if table_data:
                                create_table_in_doc(doc, table_data)
                        continue
                
                    if line:
                        if line.startswith('- ') or line.startswith('* '):
                            doc.add_paragraph(line[2:].strip(), style='List Bullet')
                        elif re.match(r'^\d+\.', line):
                            doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Bullet')
                        else:
                            doc.add_paragraph(line)
                
                    j += 1
    
    return doc

def inject_apis_table_into_section(full_text: str, api_table_md: str) -> str:
    if not api_table_md:
        return full_text
    lower = full_text.lower()
    start_markers = ["### 2.3 list of apis required", "## 2.3 list of apis required"]
    end_markers = ["## 3.0", "### 3.0", "## 3.0 process", "### 3.0 process"]
    start_idx = -1
    for m in start_markers:
        i = lower.find(m)
        if i != -1:
            start_idx = i
            break
    if start_idx == -1:
        return full_text
    # Find end of header line
    header_end = full_text.find('\n', start_idx)
    if header_end == -1:
        header_end = start_idx
    end_idx = len(full_text)
    for m in end_markers:
        j = lower.find(m, header_end + 1)
        if j != -1:
            end_idx = min(end_idx, j)
    header_text = full_text[start_idx:header_end]
    original_body = full_text[header_end:end_idx]
    # Remove default "No specific APIs" lines from original body
    cleaned_body = "\n".join([
        ln for ln in original_body.split('\n')
        if 'no specific apis' not in ln.lower()
    ]).strip()
    new_section = header_text + "\n\n" + api_table_md + ("\n\n" + cleaned_body if cleaned_body else "") + "\n"
    return full_text[:start_idx] + new_section + full_text[end_idx:]

st.title("Business Requirements Document Generator")

st.subheader("AI Model Selection")
api_provider = st.radio("Select API Provider:", ["OpenAI", "Groq", "AzureOpenAI"])

if api_provider == "OpenAI":
    api_key = st.text_input("Enter your OpenAI API Key:", type="password")
elif api_provider == "AzureOpenAI":
    api_key = st.text_input("Enter your Azure OpenAI API Key:", type="password")
    azure_endpoint = st.text_input("Enter your Azure OpenAI Endpoint:", 
                                   placeholder="https://your-resource.openai.azure.com/")
    azure_deployment = st.text_input("Enter your Azure Deployment Name:", 
                                     placeholder="gpt-35-turbo")
    api_version = st.text_input("Enter API Version (optional):", 
                                value="2025-01-01-preview",
                                placeholder="2025-01-01-preview")
else:
    api_key = st.text_input("Enter your Groq API Key:", type="password")

st.subheader("Document Logo")

logo_file = st.file_uploader("Upload Company Logo (optional):", type=['png', 'jpg', 'jpeg'])
logo_data = None
if logo_file:
    logo_data = logo_file.getvalue()
    st.success("Logo uploaded successfully!")

st.subheader("Upload Requirements Documents")
uploaded_files = st.file_uploader(
    "Choose files", 
    type=['txt', 'docx', 'pdf', 'xlsx', 'xls', 'msg'],
    accept_multiple_files=True
)

st.subheader("Or Enter Requirements Manually")
manual_requirements = st.text_area(
    "Paste your requirements here:",
    height=200,
    placeholder="Enter your business requirements, user stories, or project specifications here..."
)

if st.button("Generate BRD", type="primary"):
    if not api_key:
        st.error("Please enter your API key!")
    elif not uploaded_files and not manual_requirements.strip():
        st.error("Please upload files or enter requirements manually!")
    else:
        try:
            with st.spinner("Initializing AI chains"):
                chains = initialize_sequential_chains(api_provider=api_provider,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            azure_deployment=azure_deployment,
            api_version=api_version)
            
            all_requirements = []
            
            if manual_requirements.strip():
                all_requirements.append("=== MANUAL REQUIREMENTS ===")
                all_requirements.append(manual_requirements.strip())
                all_requirements.append("="*50)
            
            if uploaded_files:
                st.info(f"Processing {len(uploaded_files)} uploaded files...")
                
                for uploaded_file in uploaded_files:
                    file_extension = uploaded_file.name.split('.')[-1].lower()
                    
                    try:
                        st.write(f"Processing: {uploaded_file.name}")
                        
                        if file_extension == 'txt':
                            content = str(uploaded_file.read(), "utf-8")
                        elif file_extension == 'docx':
                            content = extract_content_from_docx(uploaded_file)
                        elif file_extension == 'pdf':
                            content = extract_content_from_pdf(uploaded_file)
                        elif file_extension in ['xlsx', 'xls']:
                            content = extract_content_from_excel(uploaded_file)
                        elif file_extension == 'msg':
                            content = extract_content_from_msg(uploaded_file)
                        else:
                            st.warning(f"⚠Unsupported file type: {file_extension}")
                            continue
                        
                        if content.strip():
                            all_requirements.append(f"=== FILE: {uploaded_file.name} ===")
                            all_requirements.append(content.strip())
                            all_requirements.append("="*50)
                            st.success(f"Successfully processed: {uploaded_file.name}")
                        else:
                            st.warning(f"No content extracted from: {uploaded_file.name}")
                            
                    except Exception as e:
                        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                        continue
            
            if not all_requirements:
                st.error("No valid content found in uploaded files!")
                st.stop()
            
            combined_requirements = "\n\n".join(all_requirements)
            
            content_size = estimate_content_size(combined_requirements)
            st.info(f"Total content size: {content_size:,} characters")
            
            st.subheader("AI Processing Progress")
            
            with st.spinner("Generating comprehensive BRD using sequential processing..."):
                brd_content = generate_brd_sequentially(chains, combined_requirements)
            
            if brd_content:
                st.success("BRD generated successfully!")
                
                st.subheader("Generated BRD Content")
                
                with st.expander("Preview Generated BRD", expanded=False):
                    st.markdown(brd_content)
                
                st.subheader("Download Options")
                
                try:
                    with st.spinner("Creating Word document..."):
                        doc = create_word_document(brd_content, logo_data)
                        
                        doc_buffer = BytesIO()
                        doc.save(doc_buffer)
                        doc_buffer.seek(0)
                        
                        st.download_button(
                            label="Download BRD (Word Document)",
                            data=doc_buffer.getvalue(),
                            file_name="Business_Requirements_Document.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("Word document ready for download!")
                        
                except Exception as e:
                    st.error(f"Error creating Word document: {str(e)}")
                    st.info("You can still copy the content above manually.")
                
                try:
                    st.download_button(
                        label="Download BRD (Markdown)",
                        data=brd_content,
                        file_name="Business_Requirements_Document.md",
                        mime="text/markdown"
                    )
                except Exception as e:
                    st.error(f"Error creating markdown download: {str(e)}")
                
            else:
                st.error("Failed to generate BRD content!")
                
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            st.info("Try reducing the input size or check your API key.")
